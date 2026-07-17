"""
File browser router.
All paths are validated against the user's allowed mounted_paths.
No path traversal is possible — realpath() is asserted within whitelist.
"""
import os
import mimetypes
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel

from app.core.deps import get_current_user
from app.database import get_db
from app.models.user import User, UserSettings
from app.config import settings

router = APIRouter(prefix="/api/v1/files", tags=["files"])

ALLOWED_EXTENSIONS = {
    "pdf", "docx", "doc", "pptx", "ppt", "xlsx", "xls",
    "md", "txt", "csv", "json",
    "jpg", "jpeg", "png", "gif", "webp", "svg",
    "mp4", "webm", "mov", "mp3", "wav",
}


def _allowed_roots(user_id, db: Session) -> list[str]:
    us = db.query(UserSettings).filter(UserSettings.user_id == user_id).first()
    paths = us.mounted_paths if us and us.mounted_paths else [settings.FILES_MOUNT_ROOT]
    return [str(Path(p).resolve()) for p in paths]


def _safe_path(requested: str, user_id, db: Session) -> Path:
    """Resolve path and assert it's within an allowed root."""
    roots = _allowed_roots(user_id, db)
    resolved = Path(requested).resolve()
    for root in roots:
        try:
            resolved.relative_to(root)
            return resolved
        except ValueError:
            continue
    raise HTTPException(status_code=403, detail="Path not within allowed directories")


@router.get("")
def list_directory(
    path: str = Query(default="/userfiles"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    safe = _safe_path(path, current_user.id, db)
    if not safe.exists():
        raise HTTPException(status_code=404, detail="Path not found")
    if not safe.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")

    entries = []
    try:
        all_entries = []
        for entry in safe.iterdir():
            try:
                # use lstat so broken symlinks don't raise; is_dir follows real dirs
                lstat = entry.stat(follow_symlinks=False)
                is_dir = entry.is_dir()
                all_entries.append((entry, lstat, is_dir))
            except (FileNotFoundError, PermissionError, OSError):
                continue  # skip broken symlinks / Windows junctions

        for entry, lstat, is_dir in sorted(all_entries, key=lambda t: (not t[2], t[0].name.lower())):
            entries.append({
                "name": entry.name,
                "path": str(entry),
                "is_dir": is_dir,
                "size": lstat.st_size if not is_dir else None,
                "modified": lstat.st_mtime,
                "extension": entry.suffix.lstrip(".").lower() if entry.is_file() else None,
            })
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")

    # Only expose parent if it's still within an allowed root
    roots = _allowed_roots(current_user.id, db)
    parent_path = safe.parent
    parent_str = None
    if str(parent_path) != str(safe):
        for root in roots:
            try:
                parent_path.relative_to(root)
                parent_str = str(parent_path)
                break
            except ValueError:
                continue

    return {
        "path": str(safe),
        "parent": parent_str,
        "entries": entries,
    }


@router.get("/info")
def file_info(path: str = Query(...), current_user: User = Depends(get_current_user),
              db: Session = Depends(get_db)):
    safe = _safe_path(path, current_user.id, db)
    if not safe.exists():
        raise HTTPException(status_code=404, detail="File not found")
    stat = safe.stat()
    mime, _ = mimetypes.guess_type(str(safe))
    return {
        "name": safe.name,
        "path": str(safe),
        "size": stat.st_size,
        "modified": stat.st_mtime,
        "mime_type": mime,
        "extension": safe.suffix.lstrip(".").lower(),
    }


@router.get("/preview")
def preview_file(path: str = Query(...), current_user: User = Depends(get_current_user),
                 db: Session = Depends(get_db)):
    safe = _safe_path(path, current_user.id, db)
    if not safe.exists() or not safe.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    ext = safe.suffix.lstrip(".").lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail="File type not supported for preview")

    # Text files — return raw content
    if ext in {"md", "txt", "csv", "json"}:
        try:
            content = safe.read_text(encoding="utf-8", errors="replace")
            return {"type": "text", "content": content, "extension": ext}
        except Exception:
            raise HTTPException(status_code=500, detail="Could not read file")

    # PDF — return page count info (actual rendering done client-side via react-pdf)
    if ext == "pdf":
        try:
            import fitz
            doc = fitz.open(str(safe))
            page_count = len(doc)
            doc.close()
            return {"type": "pdf", "page_count": page_count, "path": path}
        except Exception:
            return {"type": "pdf", "page_count": None, "path": path}

    # DOCX — extract text
    if ext in {"docx", "doc"}:
        try:
            import docx
            doc = docx.Document(str(safe))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {"type": "docx", "content": text}
        except Exception:
            raise HTTPException(status_code=500, detail="Could not parse document")

    # Image/video/audio — stream directly
    mime, _ = mimetypes.guess_type(str(safe))
    return FileResponse(str(safe), media_type=mime or "application/octet-stream")


@router.get("/download")
def download_file(path: str = Query(...), current_user: User = Depends(get_current_user),
                  db: Session = Depends(get_db)):
    safe = _safe_path(path, current_user.id, db)
    if not safe.exists() or not safe.is_file():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(str(safe), filename=safe.name,
                        headers={"Content-Disposition": f'attachment; filename="{safe.name}"'})


class ConvertRequest(BaseModel):
    path: str
    to: str  # 'pdf' | 'docx' | 'txt'


@router.post("/convert")
def convert_file(body: ConvertRequest, current_user: User = Depends(get_current_user),
                 db: Session = Depends(get_db)):
    safe = _safe_path(body.path, current_user.id, db)
    if not safe.exists():
        raise HTTPException(status_code=404, detail="File not found")
    # TODO: invoke LibreOffice/Pandoc for conversion
    # subprocess.run(["libreoffice", "--headless", "--convert-to", body.to, str(safe)], ...)
    raise HTTPException(status_code=501, detail="Conversion coming soon — LibreOffice required in container")


class AnalyzeRequest(BaseModel):
    path: str
    question: Optional[str] = None
    provider: str = "gemini"


@router.post("/analyze")
async def analyze_file(body: AnalyzeRequest, current_user: User = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    safe = _safe_path(body.path, current_user.id, db)
    if not safe.exists() or not safe.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    ext = safe.suffix.lstrip(".").lower()
    content = ""

    if ext == "pdf":
        import fitz
        doc = fitz.open(str(safe))
        content = "\n\n".join(page.get_text() for page in doc)[:15000]
        doc.close()
    elif ext in {"docx", "doc"}:
        import docx
        doc = docx.Document(str(safe))
        content = "\n".join(p.text for p in doc.paragraphs)[:15000]
    elif ext in {"md", "txt"}:
        content = safe.read_text(encoding="utf-8", errors="replace")[:15000]
    else:
        raise HTTPException(status_code=415, detail="File type not supported for analysis")

    from app.core.ai.provider import chat, AIMessage
    from app.models.user import UserSettings

    us = db.query(UserSettings).filter(UserSettings.user_id == current_user.id).first()
    user_providers = us.ai_providers if us else {}

    system = "You are an expert document analyst. Be concise and precise."
    prompt = body.question or "Summarize this document and extract the 5 most important points."
    messages = [AIMessage(role="user", content=f"Document content:\n\n{content}\n\n---\n\n{prompt}")]

    try:
        response = await chat(messages, provider=body.provider, user_ai_providers=user_providers,
                              system_prompt=system)
        return {"analysis": response.content, "provider": response.provider, "model": response.model}
    except Exception as e:
        raise HTTPException(status_code=500, detail="AI analysis failed — check your API key")
